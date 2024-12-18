import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
from typing import TypeVar

Document_Object = TypeVar("Document_Object")
Run_Object = TypeVar("Run_Object")
Paragraph_Object = TypeVar("Paragraph_Object")

def search_in_pdf(file_path: str, search_term: str) -> list[dict[int, str]]:
    """
    Searches for a term in a PDF file.
    Args:
    file_path (str): The path to the PDF file.
    search_term (str): The term to search for.
    Returns:
    list[dict[int, str]]: A list of dictionaries containing the page number, the text and the index
    """
    results: list[dict[int, str]] = []
    with fitz.open(file_path) as doc:
        for page_num in range(len(doc)):
            page: str = doc[page_num]
            text: str = page.get_text("text").lower()
            index: int = text.find(search_term)
            
            while index != -1:
                start: int = max(0, index - 300)
                end: int = index + len(search_term) + 300
                context: str = text[start:end]
                results.append({
                    'page': page_num + 1,
                    'context': context,
                    'index': index - start
                })
                index = text.find(search_term, index + len(search_term))
    return results

def search_in_folder(folder_path: str, search_term: str) -> dict[list[dict[int, str]]]:
    """
    Searches for a term in all PDF files in a folder.
    Args:
    folder_path (str): The path to the folder containing the PDF files.
    search_term (str): The term to search for.
    Returns:
    dict[list[dict[int, str]]]: A dictionary with the file name as key and the
                                search results as value.
    """
    pdf_files: list[str] = [f for f in os.listdir(folder_path) if f.endswith('.pdf')]
    all_results: dict[list[dict[int, str]]] = {}
    for pdf_file in pdf_files:
        file_path = os.path.join(folder_path, pdf_file)
        results = search_in_pdf(file_path, search_term)
        if results:
            all_results[pdf_file] = results
    return all_results

def highlight_search_term(paragraph: Paragraph_Object, search_term: str) -> None:
    """
    Highlights the search term in a paragraph.
    Args:
    paragraph (Paragraph_Object): The paragraph to highlight the search term in.
    search_term (str): The term to highlight.
    """
    run: Run_Object = paragraph.add_run(search_term)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)

def sanitize_text(text: str) -> str:
    """
    Sanitizes the text by removing special characters (except \n and \t) and converting to lowercase.
    Args:
    text (str): The text to sanitize.
    Returns:
    str: The sanitized text.
    """
    return ''.join(c for c in text if c.isprintable() or c in '\n\t')

def write_results_to_docx(all_results: dict[list[dict[int, str]]], search_term: str, output_path: str) -> None:
    """
    Writes the search results to a docx file.
    Args:
    all_results (dict[list[dict[int, str]]]): The search results.
    search_term (str): The term that was searched for.
    output_path (str): The path to the output docx file.
    """
    doc: Document_Object = Document()
    doc.add_heading(f'Search Results for "{search_term}"', level=1)

    for pdf_file, results in all_results.items():
        doc.add_heading(pdf_file, level=2)
        for result in results:
            doc.add_heading(f'Page {result["page"]}', level=3)
            context: str = sanitize_text(result['context'])
            index = result['index']

            # Split the context around the search term
            before = context[:index]
            after = context[index + len(search_term):]

            paragraph: Paragraph_Object = doc.add_paragraph(sanitize_text(before))
            highlight_search_term(paragraph, search_term)
            paragraph.add_run(sanitize_text(after))

    doc.save(f"{output_path}/occurrence_of_{search_term}.docx")

if __name__ == "__main__":
    folder_path: str = input("Enter the folder path containing PDF files: ")
    search_term: str = input("Enter the search term: ")
    output_path: str = input("Enter the output DOCX file path: ")

    all_results: dict[list[dict[int, str]]] = search_in_folder(folder_path, search_term)
    write_results_to_docx(all_results, search_term, output_path)
    print(f'Results written to {output_path}')